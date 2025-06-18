from docx import Document
from docx.shared import Pt

def export_structured_resume(data, filename="tailored_resume.docx"):
    doc = Document()

    def add_heading(text):
        doc.add_heading(text, level=1)

    def add_bullet(text):
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(text)
        run.font.size = Pt(10.5)

    def add_section(title, lines):
        doc.add_heading(title, level=2)
        for line in lines:
            if isinstance(line, str):
                add_bullet(line)
            elif isinstance(line, dict):
                doc.add_paragraph(line.get("title", ""), style="Heading 3")
                for bullet in line.get("bullets", []):
                    add_bullet(bullet)

    # Header
    header = doc.add_paragraph()
    header.add_run(data.get("name", "")).bold = True
    header.add_run(" | ").bold = True
    header.add_run(data.get("email", ""))
    header.add_run(" | ").bold = True
    header.add_run(data.get("linkedin", ""))

    # Sections
    if "education" in data:
        add_section("Education", data["education"])

    if "experience" in data:
        add_section("Experience", data["experience"])

    if "projects" in data:
        add_section("Projects", data["projects"])

    if "skills" in data:
        add_section("Skills", data["skills"])

    doc.save(filename)
