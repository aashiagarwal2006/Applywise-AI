from docx import Document

def export_resume_to_docx(text, filename="tailored_resume.docx"):
    doc = Document()
    doc.add_heading('Tailored Resume', 0)

    for line in text.split("\n"):
        if line.strip() == "":
            continue
        elif line.strip().endswith(":"):
            doc.add_heading(line.strip()[:-1], level=1)
        elif line.strip().startswith("-"):
            doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.save(filename)
